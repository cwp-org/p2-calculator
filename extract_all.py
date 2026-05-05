import openpyxl
from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

# ─────────────────────────────────────────────
# PART 1 — EXCEL FILE
# ─────────────────────────────────────────────
EXCEL_PATH = "/home/user/p2-calculator/Existing Calculator Summary_05052026.xlsx"
TARGET_SHEETS = ["Plastics Ramblings", "Online Calculator Edit Tracking"]

wb = openpyxl.load_workbook(EXCEL_PATH, data_only=True)
print("=" * 80)
print("EXCEL FILE:", EXCEL_PATH)
print("Available sheets:", wb.sheetnames)
print("=" * 80)

for sheet_name in TARGET_SHEETS:
    if sheet_name not in wb.sheetnames:
        print(f"\n[WARNING] Sheet '{sheet_name}' NOT FOUND in workbook.\n")
        continue

    ws = wb[sheet_name]
    print(f"\n{'=' * 80}")
    print(f"SHEET: {sheet_name}")
    print(f"Dimensions: {ws.dimensions}  |  Max row: {ws.max_row}  |  Max col: {ws.max_column}")
    print(f"{'=' * 80}\n")

    # Print column headers row (A, B, C …) for orientation
    col_letters = [openpyxl.utils.get_column_letter(c) for c in range(1, ws.max_column + 1)]
    print("COL LETTERS:", " | ".join(col_letters))
    print()

    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=ws.max_row,
                                               min_col=1, max_col=ws.max_column), start=1):
        # Skip completely empty rows
        values = [cell.value for cell in row]
        if all(v is None for v in values):
            continue

        print(f"  ROW {row_idx:>4}:", end="")
        for col_idx, cell in enumerate(row, start=1):
            col_letter = openpyxl.utils.get_column_letter(col_idx)
            val = cell.value
            if val is not None:
                print(f"\n           [{col_letter}{row_idx}] {repr(val)}", end="")
        print()  # newline after each row

    print(f"\n--- END OF SHEET: {sheet_name} ---\n")

# ─────────────────────────────────────────────
# PART 2 — WORD DOCUMENT
# ─────────────────────────────────────────────
DOCX_PATH = "/home/user/p2-calculator/Plastics references with links for calculator.docx"

print("\n" + "=" * 80)
print("WORD DOCUMENT:", DOCX_PATH)
print("=" * 80 + "\n")

doc = Document(DOCX_PATH)

# Helper: extract hyperlinks from a paragraph's XML
# Hyperlinks live in <w:hyperlink> elements; the relationship target is the URL.
def get_para_content(para, doc):
    """Return list of (text, url_or_None) tuples for a paragraph."""
    # Build a map of relationship id -> url from the document's part
    rel_map = {}
    for rel_id, rel in para.part.rels.items():
        if "hyperlink" in rel.reltype:
            rel_map[rel_id] = rel._target

    segments = []
    for child in para._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "hyperlink":
            # Extract r:id attribute
            r_id = child.get(qn("r:id"))
            url = rel_map.get(r_id, None)
            # Gather text inside the hyperlink
            texts = [t.text for t in child.iter(qn("w:t")) if t.text]
            link_text = "".join(texts)
            segments.append((link_text, url))
        elif tag == "r":  # plain run
            texts = [t.text for t in child.iter(qn("w:t")) if t.text]
            run_text = "".join(texts)
            if run_text:
                segments.append((run_text, None))
    return segments

para_num = 0
for para in doc.paragraphs:
    full_text = para.text.strip()
    segments = get_para_content(para, doc)

    # Skip truly empty paragraphs with no hyperlinks
    has_content = bool(full_text) or any(url for _, url in segments)
    if not has_content:
        continue

    para_num += 1
    style = para.style.name if para.style else "Normal"
    print(f"[PARA {para_num:>3}] Style={style!r}")

    if segments:
        for seg_text, seg_url in segments:
            if seg_url:
                print(f"          LINK TEXT : {seg_text!r}")
                print(f"          LINK URL  : {seg_url}")
            elif seg_text.strip():
                print(f"          TEXT      : {seg_text!r}")
    else:
        print(f"          TEXT      : {full_text!r}")
    print()

# Also check tables in the Word doc
if doc.tables:
    print("\n" + "=" * 80)
    print("TABLES IN WORD DOCUMENT")
    print("=" * 80)
    for t_idx, table in enumerate(doc.tables, start=1):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows, start=1):
            row_vals = [cell.text.strip() for cell in row.cells]
            if any(row_vals):
                print(f"  Row {r_idx:>3}: {' | '.join(row_vals)}")
else:
    print("\n[No tables found in Word document]")

print("\n" + "=" * 80)
print("EXTRACTION COMPLETE")
print("=" * 80)
